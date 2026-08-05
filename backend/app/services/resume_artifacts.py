from __future__ import annotations

import io
import re
from dataclasses import dataclass
from typing import Any, Literal
from xml.sax.saxutils import escape

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

from .. import models

ArtifactFormat = Literal["pdf", "docx"]


@dataclass(frozen=True)
class RenderedResumeArtifact:
    content: bytes
    media_type: str
    filename: str


@dataclass(frozen=True)
class ResumeDocument:
    name: str
    contact_line: str
    target_title: str
    summary: list[str]
    highlights: list[str]
    skills: list[str]
    sections: list[tuple[str, str]]


SECTION_ORDER = (
    "experience",
    "projects",
    "education",
    "certifications",
    "awards",
    "publications",
    "volunteer",
    "languages",
    "interests",
)
INVALID_DOCUMENT_CHARACTERS = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f]")


def _text(value: Any) -> str:
    return INVALID_DOCUMENT_CHARACTERS.sub("", str(value or "")).strip()


def _item_texts(value: Any, *, limit: int) -> list[str]:
    if not isinstance(value, list):
        return []
    result: list[str] = []
    for item in value[:limit]:
        text = _text(item.get("text") if isinstance(item, dict) else item)
        if text:
            result.append(text)
    return result


def _skill_texts(value: Any) -> list[str]:
    if not isinstance(value, list):
        return []
    return list(dict.fromkeys(_text(skill) for skill in value if _text(skill)))[:80]


def _section_title(key: str) -> str:
    return key.replace("_", " ").replace("-", " ").title()


def build_resume_document(
    version: models.ResumeVersion,
    resume: models.Resume,
) -> ResumeDocument:
    content = version.structured_content if isinstance(version.structured_content, dict) else {}
    stored_sections = content.get("sections")
    sections = stored_sections if isinstance(stored_sections, dict) else (resume.sections or {})
    stored_contact = content.get("contact_info")
    contact = stored_contact if isinstance(stored_contact, dict) else (resume.contact_info or {})

    name = _text(contact.get("name")) or "Professional Candidate"
    contact_line = " | ".join(
        value
        for value in (
            _text(contact.get("email")),
            _text(contact.get("phone")),
            _text(contact.get("linkedin")),
            _text(contact.get("github")),
        )
        if value
    )
    generated_summary = _item_texts(content.get("summary_items"), limit=3)
    generated_highlights = _item_texts(content.get("bullets"), limit=16)
    original_summary = _text(sections.get("summary")) if isinstance(sections, dict) else ""
    summary = generated_summary or ([original_summary] if original_summary else [])
    skills = _skill_texts(content.get("skills")) or _skill_texts(resume.skills or [])

    ordered_sections: list[tuple[str, str]] = []
    if isinstance(sections, dict):
        keys = [key for key in SECTION_ORDER if _text(sections.get(key))]
        keys.extend(
            key
            for key, value in sections.items()
            if key not in {*SECTION_ORDER, "summary", "skills", "other"} and _text(value)
        )
        if not keys and _text(sections.get("other")):
            keys.append("other")
        for key in keys:
            title = "Additional Information" if key == "other" else _section_title(key)
            ordered_sections.append((title, _text(sections[key])))

    return ResumeDocument(
        name=name,
        contact_line=contact_line,
        target_title=_text(content.get("target_job_title")),
        summary=summary,
        highlights=generated_highlights,
        skills=skills,
        sections=ordered_sections,
    )


def _filename(version: models.ResumeVersion, artifact_format: ArtifactFormat) -> str:
    stem = re.sub(r"[^A-Za-z0-9]+", "-", _text(version.label)).strip("-").lower()[:72]
    stem = stem or "tailored-resume"
    return f"{stem}-v{version.version_number}.{artifact_format}"


def _add_docx_text(document: Document, text: str) -> None:
    for raw_line in text.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        is_bullet = line.startswith(("- ", "* ", "• "))
        paragraph = document.add_paragraph(style="List Bullet" if is_bullet else None)
        paragraph.add_run(line[2:] if is_bullet else line)


def _render_docx(document_data: ResumeDocument) -> bytes:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)

    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(9.5)
    normal.paragraph_format.space_after = Pt(3)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run(document_data.name)
    title_run.bold = True
    title_run.font.size = Pt(20)
    if document_data.target_title:
        target = document.add_paragraph()
        target.alignment = WD_ALIGN_PARAGRAPH.CENTER
        target_run = target.add_run(document_data.target_title)
        target_run.bold = True
        target_run.font.size = Pt(10.5)
    if document_data.contact_line:
        contact = document.add_paragraph(document_data.contact_line)
        contact.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def heading(text: str) -> None:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(7)
        paragraph.paragraph_format.space_after = Pt(3)
        run = paragraph.add_run(text.upper())
        run.bold = True
        run.font.size = Pt(11)

    if document_data.summary:
        heading("Professional Summary")
        for item in document_data.summary:
            document.add_paragraph(item)
    if document_data.skills:
        heading("Skills")
        document.add_paragraph(" • ".join(document_data.skills))
    if document_data.highlights:
        heading("Relevant Experience Highlights")
        for item in document_data.highlights:
            document.add_paragraph(item, style="List Bullet")
    for title_text, body in document_data.sections:
        heading(title_text)
        _add_docx_text(document, body)

    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def _render_pdf(document_data: ResumeDocument) -> bytes:
    output = io.BytesIO()
    document = SimpleDocTemplate(
        output,
        pagesize=A4,
        rightMargin=0.58 * inch,
        leftMargin=0.58 * inch,
        topMargin=0.5 * inch,
        bottomMargin=0.5 * inch,
        title=f"{document_data.name} Resume",
        author=document_data.name,
    )
    base = getSampleStyleSheet()
    styles = {
        "name": ParagraphStyle(
            "ResumeName",
            parent=base["Title"],
            fontName="Helvetica-Bold",
            fontSize=20,
            leading=23,
            alignment=TA_CENTER,
            textColor=colors.HexColor("#111827"),
            spaceAfter=3,
        ),
        "target": ParagraphStyle(
            "ResumeTarget",
            parent=base["Normal"],
            fontName="Helvetica-Bold",
            fontSize=10,
            leading=13,
            alignment=TA_CENTER,
            textColor=colors.HexColor("#374151"),
            spaceAfter=2,
        ),
        "contact": ParagraphStyle(
            "ResumeContact",
            parent=base["Normal"],
            fontName="Helvetica",
            fontSize=8.5,
            leading=11,
            alignment=TA_CENTER,
            textColor=colors.HexColor("#4B5563"),
            spaceAfter=8,
        ),
        "heading": ParagraphStyle(
            "ResumeHeading",
            parent=base["Heading2"],
            fontName="Helvetica-Bold",
            fontSize=10.5,
            leading=13,
            textColor=colors.HexColor("#111827"),
            borderColor=colors.HexColor("#9CA3AF"),
            borderWidth=0,
            borderPadding=0,
            spaceBefore=7,
            spaceAfter=3,
        ),
        "body": ParagraphStyle(
            "ResumeBody",
            parent=base["BodyText"],
            fontName="Helvetica",
            fontSize=9,
            leading=12,
            textColor=colors.HexColor("#1F2937"),
            spaceAfter=3,
        ),
        "bullet": ParagraphStyle(
            "ResumeBullet",
            parent=base["BodyText"],
            fontName="Helvetica",
            fontSize=9,
            leading=12,
            leftIndent=11,
            firstLineIndent=-7,
            textColor=colors.HexColor("#1F2937"),
            spaceAfter=2,
        ),
    }
    story: list[Any] = [Paragraph(escape(document_data.name), styles["name"])]
    if document_data.target_title:
        story.append(Paragraph(escape(document_data.target_title), styles["target"]))
    if document_data.contact_line:
        story.append(Paragraph(escape(document_data.contact_line), styles["contact"]))
    else:
        story.append(Spacer(1, 5))

    def heading(text: str) -> None:
        story.append(Paragraph(escape(text.upper()), styles["heading"]))

    if document_data.summary:
        heading("Professional Summary")
        for item in document_data.summary:
            story.append(Paragraph(escape(item), styles["body"]))
    if document_data.skills:
        heading("Skills")
        story.append(Paragraph(escape(" | ".join(document_data.skills)), styles["body"]))
    if document_data.highlights:
        heading("Relevant Experience Highlights")
        for item in document_data.highlights:
            story.append(Paragraph(f"- {escape(item)}", styles["bullet"]))
    for title_text, body in document_data.sections:
        heading(title_text)
        for raw_line in body.splitlines():
            line = raw_line.strip()
            if not line:
                continue
            is_bullet = line.startswith(("- ", "* ", "• "))
            text = line[2:] if is_bullet else line
            story.append(
                Paragraph(
                    f"- {escape(text)}" if is_bullet else escape(text),
                    styles["bullet"] if is_bullet else styles["body"],
                )
            )

    document.build(story)
    return output.getvalue()


def render_resume_version(
    version: models.ResumeVersion,
    resume: models.Resume,
    artifact_format: ArtifactFormat,
) -> RenderedResumeArtifact:
    document = build_resume_document(version, resume)
    if artifact_format == "docx":
        return RenderedResumeArtifact(
            content=_render_docx(document),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            filename=_filename(version, artifact_format),
        )
    return RenderedResumeArtifact(
        content=_render_pdf(document),
        media_type="application/pdf",
        filename=_filename(version, artifact_format),
    )
